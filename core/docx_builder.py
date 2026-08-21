# -*- coding: utf-8 -*-
"""
docx 조립 모듈. Call2~6 결과와 chart_generator가 만든 차트 이미지(BytesIO)를 받아
GUIDE_v2.1 섹션 6 스타일 규칙에 따라 최종 .docx를 BytesIO로 반환한다.
(Node.js 버전 build.js / 프로토타입 build_full.py의 python-docx 이식 검증을 그대로 재사용)
"""
import io
import re
from docx import Document
from docx.shared import Twips, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY = "1F4E79"
SECONDARY = "2E75B6"
LIGHT = "EBF3FA"

KOPUB_BOLD = "KoPub돋움체 Bold"
KOPUB_MEDIUM = "KoPub돋움체 Medium"
KOPUB_LIGHT = "KoPub돋움체 Light"

PAGE_W, PAGE_H = 11906, 16838
MARGIN = dict(top=1134, bottom=1134, left=1417, right=1134)
TABLE_WIDTH = PAGE_W - MARGIN["left"] - MARGIN["right"]


def _font(run, name=KOPUB_MEDIUM, size=10.5, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)


_SENTENCE_SPLIT_RE = re.compile(r'(?<=[.!?])\s+(?=[가-힣A-Za-z0-9])')
_URL_RE = re.compile(r"https?://[^\s]+")


def _split_paragraphs(text: str, sentences_per_para: int = 2) -> list:
    sentences = [s.strip() for s in _SENTENCE_SPLIT_RE.split(text.strip()) if s.strip()]
    if len(sentences) <= sentences_per_para:
        return [text.strip()]
    return [
        " ".join(sentences[i:i + sentences_per_para])
        for i in range(0, len(sentences), sentences_per_para)
    ]


def _normalize_cell(value) -> str:
    return re.sub(r"\s+", "", str(value or "")).lower()


def _shorten_cell(value, limit: int = 220) -> str:
    text = str(value or "").replace("\n", " ").strip()
    text = re.sub(r"\s+", " ", text)
    text = _URL_RE.sub(lambda m: _shorten_url(m.group(0)), text)
    if len(text) <= limit:
        return text
    return text[:limit].rstrip() + "..."


def _shorten_url(url: str) -> str:
    body = url.split("://", 1)[-1]
    domain = body.split("/", 1)[0]
    return f"{domain}/..."


def _clean_rows(headers, rows):
    cleaned = []
    header_norm = [_normalize_cell(h) for h in headers]
    for row in rows or []:
        values = list(row) if isinstance(row, (list, tuple)) else [row]
        values = values[:len(headers)] + [""] * max(0, len(headers) - len(values))
        row_norm = [_normalize_cell(v) for v in values]
        if row_norm == header_norm:
            continue
        if _looks_like_schema_example(row_norm):
            continue
        cleaned.append([_shorten_cell(v) for v in values])
    return cleaned


def _looks_like_schema_example(row_norm) -> bool:
    examples = {
        ("유형", "매체", "성과"),
        ("분류", "단계", "r&d성장도", "시사점"),
        ("분류", "단계", "rd성장도", "시사점"),
        ("출원인", "시기별변화", "전환패턴"),
        ("지표", "설명"),
        ("분류", "범위"),
        ("국가", "정책동향"),
        ("기업명", "역할", "동향"),
        ("영역", "r&d성장", "특허성숙", "전략"),
        ("영역", "rd성장", "특허성숙", "전략"),
    }
    compact = tuple(row_norm)
    return compact in examples


def _classification_label(row) -> tuple[str | None, str | None]:
    if not row:
        return None, None
    text = str(row[0]).strip()
    match = re.match(r"^([A-D])(?:[.\s]|$)", text)
    if not match:
        return None, text or None
    code = match.group(1)
    return code, text


def _count_candidate_rows(rows):
    converted = []
    for row in rows or []:
        values = list(row) if isinstance(row, (list, tuple)) else [row]
        if len(values) >= 2:
            values[1] = _as_candidate_range(values[1])
        converted.append(values)
    return converted


def _as_candidate_range(value) -> str:
    text = str(value or "").strip()
    if "~" in text or "범위" in text:
        return text
    match = re.search(r"(\d[\d,]*)", text)
    if not match:
        return text
    number = int(match.group(1).replace(",", ""))
    low = max(1, round(number * 0.8))
    high = max(low, round(number * 1.2))
    return f"약 {low:,}~{high:,}건(E, 후보)"


def _shade(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _width(cell, width_dxa):
    cell.width = Twips(width_dxa)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_dxa)); tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def _page_number_field(paragraph, field_type="PAGE"):
    run = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    t = OxmlElement('w:instrText'); t.set(qn('xml:space'), 'preserve'); t.text = field_type
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    run._r.append(b); run._r.append(t); run._r.append(e)
    _font(run, name=KOPUB_LIGHT, size=8, color="888888")


class DocxBuilder:
    def __init__(self, tech_name, purpose, scenario_label):
        self.doc = Document()
        self.tech_name = tech_name
        self.purpose = purpose
        self.scenario_label = scenario_label
        self._setup_page()
        self._setup_header_footer()

    def _setup_page(self):
        s = self.doc.sections[0]
        s.page_width = Twips(PAGE_W); s.page_height = Twips(PAGE_H)
        s.top_margin = Twips(MARGIN["top"]); s.bottom_margin = Twips(MARGIN["bottom"])
        s.left_margin = Twips(MARGIN["left"]); s.right_margin = Twips(MARGIN["right"])
        self.section = s

    def _setup_header_footer(self):
        hp = self.section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _font(hp.add_run(f"{self.tech_name} | 기술동향 분석 보고서"), name=KOPUB_LIGHT, size=8, color="888888")

        fp = self.section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(fp.add_run(f"{self.tech_name} | {self.purpose} | "), name=KOPUB_LIGHT, size=8, color="888888")
        _page_number_field(fp, "PAGE")
        _font(fp.add_run(" / "), name=KOPUB_LIGHT, size=8, color="888888")
        _page_number_field(fp, "NUMPAGES")

    def cover(self, date_str):
        for _ in range(6):
            self.doc.add_paragraph()
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(p.add_run(f"{self.tech_name}\n기술동향 분석 보고서"), name=KOPUB_BOLD, size=22, color=PRIMARY)
        p2 = self.doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(p2.add_run("기술동향 분석 보고서"), name=KOPUB_MEDIUM, size=14, color=SECONDARY)
        for _ in range(3):
            self.doc.add_paragraph()
        for label, val in [("작성 일자", date_str), ("분석 목적", self.purpose), ("적용 시나리오", self.scenario_label)]:
            pp = self.doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _font(pp.add_run(f"{label}: {val}"), name=KOPUB_MEDIUM, size=11)
        disc = self.doc.add_paragraph(); disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        disc.paragraph_format.space_before = Pt(30)
        _font(disc.add_run("※ 본 보고서의 정량 수치는 별도 출처 표기가 없는 한 AI 기반 추정치(E)이며,\n"
                            "실제 의사결정 이전 원 데이터 재검증이 필요합니다."), name=KOPUB_LIGHT, size=9, italic=True, color="777777")

    def h1(self, text, page_break=True):
        if page_break:
            self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        p = self.doc.add_paragraph(style="Heading 1")
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(12)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '6'); bottom.set(qn('w:color'), PRIMARY)
        pBdr.append(bottom); pPr.append(pBdr)
        _font(p.add_run(text), name=KOPUB_BOLD, size=16, color=PRIMARY)

    def h2(self, text):
        p = self.doc.add_paragraph(style="Heading 2")
        p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(8)
        _font(p.add_run(text), name=KOPUB_BOLD, size=13, color=SECONDARY)

    def h3(self, text):
        p = self.doc.add_paragraph(style="Heading 3")
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6)
        _font(p.add_run(text), name=KOPUB_MEDIUM, size=11, color="333333")

    def body(self, text):
        for chunk in _split_paragraphs(text):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12); p.paragraph_format.line_spacing = 1.25
            _font(p.add_run(chunk), name=KOPUB_MEDIUM, size=10.5)

    def bullet(self, text):
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(8)
        _font(p.add_run(text), name=KOPUB_MEDIUM, size=10.5)

    def table(self, headers, rows, ratios):
        rows = _clean_rows(headers, rows)
        total = sum(ratios)
        widths = [round(r / total * TABLE_WIDTH) for r in ratios]
        widths[-1] += TABLE_WIDTH - sum(widths)
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.autofit = False
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            _width(hdr[i], widths[i]); _shade(hdr[i], PRIMARY)
            pp = hdr[i].paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _font(pp.add_run(h), name=KOPUB_BOLD, size=9.5, color="FFFFFF")
        for ridx, row in enumerate(rows):
            cells = t.add_row().cells
            trPr = cells[0]._tc.getparent().get_or_add_trPr()
            cant_split = OxmlElement('w:cantSplit')
            trPr.append(cant_split)
            fill = LIGHT if ridx % 2 == 1 else None
            for i, val in enumerate(row):
                _width(cells[i], widths[i])
                if fill:
                    _shade(cells[i], fill)
                _font(cells[i].paragraphs[0].add_run(str(val)), name=KOPUB_MEDIUM, size=9.5)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def image(self, buf: io.BytesIO, width_cm, caption):
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        buf.seek(0)
        p.add_run().add_picture(buf, width=Cm(width_cm))
        cap = self.doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        _font(cap.add_run(caption), name=KOPUB_LIGHT, size=8.5, italic=True, color="555555")

    def note(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(10)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'FFF7E6')
        pPr.append(shd)
        pBdr = OxmlElement('w:pBdr')
        for side in ('top', 'bottom', 'left', 'right'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '4'); el.set(qn('w:space'), '4'); el.set(qn('w:color'), 'EDA100')
            pBdr.append(el)
        pPr.append(pBdr)
        _font(p.add_run(f"※ {text}"), name=KOPUB_MEDIUM, size=9.5, italic=True, color="6B4E00")

    def save_to_buffer(self) -> io.BytesIO:
        buf = io.BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        return buf


def _rows_from_links(items):
    rows = []
    for item in items or []:
        rows.append([
            item.get("source", ""),
            item.get("purpose", ""),
            item.get("query", ""),
            item.get("grade", ""),
            item.get("url", ""),
        ])
    return rows


def build_report_docx(tech_name, purpose, scenario_label, date_str, chapter_results, chart_images: dict, confirmed_context: dict | None = None) -> io.BytesIO:
    b = DocxBuilder(tech_name, purpose, scenario_label)
    b.cover(date_str)
    research_plan = (confirmed_context or {}).get("research_plan", {})

    ch1 = chapter_results["call2"]["ch1"]; ch2 = chapter_results["call2"]["ch2"]
    ch3 = chapter_results["call3"]["ch3"]; ch4 = chapter_results["call3"]["ch4"]
    ch5 = chapter_results["call4"]["ch5"]
    ch6 = chapter_results["call5"]["ch6"]; ch7 = chapter_results["call5"]["ch7"]
    ch8 = chapter_results["call6"]["ch8"]; ch9 = chapter_results["call6"]["ch9"]

    if research_plan:
        b.h1("분석 근거 및 무료 검증 계획", page_break=True)
        input_basis = research_plan.get("input_basis", {})
        b.body(
            f"본 보고서는 업로드 기술문서와 사용자 입력을 1차 근거로 생성되며, "
            f"현재 입력 근거 등급은 {input_basis.get('grade', 'D')}입니다. "
            f"{input_basis.get('reason', '')}"
        )
        b.note(
            "이 장의 검색 링크와 출처명은 검증을 위한 경로이며, 원문 수치나 검색결과를 직접 집계한 확정 근거가 아니다. "
            "외부 제출용으로 사용할 때는 검색일, 검색식, 원자료 파일, 원문 출처를 별도로 남겨야 한다."
        )
        b.h2("1. 근거 등급 기준")
        b.table(["등급", "의미", "적용 조건"], research_plan.get("evidence_grades", []), [0.7, 2, 3.4])
        b.h2("2. 특허 무료 검색/검증 링크")
        patent_rows = _rows_from_links(research_plan.get("patent_search_links"))
        if patent_rows:
            b.table(["검색원", "용도", "검색식", "등급", "링크(축약)"], patent_rows, [0.9, 1.5, 2.5, 0.6, 1.8])
        b.h2("3. 논문·시장 공개자료 검색 링크")
        public_rows = _rows_from_links(research_plan.get("scholar_search_links")) + _rows_from_links(research_plan.get("market_search_links"))
        if public_rows:
            b.table(["검색원", "용도", "검색식", "등급", "링크(축약)"], public_rows, [0.9, 1.5, 2.5, 0.6, 1.8])
        b.note("무료 공개 검색 기반 수치는 실제 검색일, 검색식, 다운로드 파일 유무에 따라 신뢰도가 달라진다. 외부 제출 전에는 검색결과 파일 또는 원문 출처로 재검증해야 한다.")

    b.h1("Ⅰ. 기술 개요 및 배경", page_break=True)
    b.h2("1. 기술 정의 및 배경"); b.body(ch1["intro"]); b.body(ch1["background"])
    b.h2("2. 주요 성능 지표 (KPI)"); b.table(["지표", "설명"], ch1["kpi"], [1, 3])
    b.h2("3. 기술 분류 체계"); b.table(["분류", "범위"], ch1["classification_desc"], [1, 3])

    b.h1("Ⅱ. 시장 환경 분석")
    b.h2("1. 시장 규모 및 성장 전망"); b.body(ch2["market_intro"])
    if ch2.get("market_confidence"):
        b.note(ch2["market_confidence"])
    b.h2("2. 주요국 정책 동향"); b.table(["국가/지역", "정책 동향"], ch2["policy"], [1, 4])
    b.h2("3. 주요 기업 동향"); b.table(["기업명", "역할", "주요 동향"], ch2["players"], [1, 1, 3.2])
    if ch2.get("market_sources"):
        b.h2("3-1. 시장자료 출처 후보 및 검증 수준")
        b.table(["출처/검색경로", "확인할 항목", "근거등급", "비고"], ch2["market_sources"], [1.4, 1.5, 0.7, 2.4])
    b.h2("4. 시장 규모 전망 출처별 비교")
    b.image(chart_images["chart5_market"], 14.5, "[Chart 5] 출처별 시장 규모·CAGR 후보 비교 (E, 차트용 대표값)")
    b.h2("5. 주요 기업 포지셔닝 맵")
    b.image(chart_images["chart6_positioning"], 13, "[Chart 6] 주요 기업 포지셔닝 후보 맵 (E)")

    b.h1("Ⅲ. 특허 정량 분석")
    b.h2("1. 분석 개요"); b.body(ch3["overview"]); b.table(["기술 분류", "특허 건수 후보 범위"], _count_candidate_rows(ch3["counts"]), [1, 2])
    strategy = ch3.get("patent_search_strategy", {})
    if strategy:
        if strategy.get("verification_note"):
            b.note(strategy["verification_note"])
        if strategy.get("search_queries"):
            b.h2("1-1. 특허 검색식 및 검증 링크")
            b.table(["검색원", "검색식", "검증링크(축약)", "근거등급"], strategy["search_queries"], [1, 2.8, 1.8, 0.8])
        if strategy.get("ipc_cpc_candidates"):
            b.h2("1-2. IPC/CPC 후보")
            b.table(["코드", "선정근거"], strategy["ipc_cpc_candidates"], [0.8, 4])
    b.h2("2. 전체 출원 동향")
    b.image(chart_images["chart4_country"], 15.5, "[Chart 4] 국가별 연도별 특허 출원 후보 추이 및 누적 비중 (E)")
    b.h2("3. 주요 출원인 후보 현황")
    b.table(["출원인 후보", "건수 동향 후보", "주요 국가", "핵심 기술 영역"], ch3["applicants"], [1, 1, 1, 2])
    b.h2("4. 주요 IPC 분류 동향"); b.table(["IPC 코드", "의미"], ch3["ipc"], [1, 3])

    b.h1("Ⅳ. 핵심 기술 개요")
    b.table(["기술 분류", "핵심 원리", "기술 계보(진화 방향)", "응용/상용화 포인트"], ch4["summary"], [1, 2, 2.4, 1.8])

    b.h1("Ⅴ. 주요 R&D 동향 분석")
    b.h2("1. R&D 분석 개요")
    b.image(chart_images["chart1_trend"], 15.5, "[Chart 1] 기술 영역별 연도별 논문 건수 후보 추이 (E)")
    for t in ch5["trends"]:
        b.bullet(t)
    b.h2("2. 기술 영역별 R&D 동향")
    classification_labels = {}
    for row in ch1["classification_desc"]:
        code, label = _classification_label(row)
        if code and label:
            classification_labels[code] = label
    for key in ["A", "B", "C", "D"]:
        if key in ch5["by_area"]:
            b.h3(classification_labels.get(key, key)); b.body(ch5["by_area"][key])
    b.h2("주요 연구 그룹 및 성과")
    b.table(["연구 그룹 유형(E)", "발표 매체", "핵심 성과 영역"], ch5["research_groups"], [1.6, 1.4, 2.4])
    b.h2("3. 핵심 키워드 트렌드")
    b.image(chart_images["chart2_keywords"], 15, "[Chart 2] 핵심 키워드 출현 빈도 후보 비교 (E)")
    b.h2("4. R&D-특허 종합 교차 분석")
    b.image(chart_images["chart3_matrix"], 13, "[Chart 3] R&D-특허 포지셔닝 후보 매트릭스 (E)")
    b.table(["기술 영역", "R&D 성장도", "특허 성숙도", "권장 전략"], ch5["cross_analysis"], [1.2, 1, 1, 2.6])

    b.h1("Ⅵ. 기술 성장 단계 분석")
    b.table(["기술 분류", "성장 단계", "R&D 성장도", "시사점"], ch6["stages"], [1.2, 1.2, 1, 2.6])
    b.h2("종합 판단"); b.body(ch6["overall"])

    b.h1("Ⅶ. 주요 출원인 후보 IP 히스토리 분석")
    b.h2("1. 상위 출원인 후보 IP 히스토리")
    b.table(["출원인 후보", "시기별 핵심 출원 영역 변화 후보(E)", "R&D-특허 전환 패턴"], ch7["history"], [1, 3.2, 2])
    b.h2("2. 의뢰 기관 보유 IP 현황 및 평가"); b.body(ch7["own_ip_note"])

    b.h1("Ⅷ. 공백기술 도출 및 IP 포트폴리오 전략")
    b.h2("1. 공백기술 도출 (3P 분석: 특허·논문·시장 교차)")
    b.table(["No.", "공백기술명", "관련 선행특허", "신규 IP 창출안", "진입 가능성"], ch8["gaps"], [0.4, 1.8, 1.8, 2.4, 1.6])
    b.h2("2. IP 재구성 전략 (4관점)"); b.table(["관점", "전략 내용"], ch8["reorg_strategy"], [1, 3])

    b.h1("Ⅸ. 결론 및 시사점")
    b.h2("1. 기술·R&D 동향 요약")
    for k in ch9["key_points"]:
        b.bullet(k)
    b.h2("2. 우선순위 행동 과제"); b.table(["시점", "과제명", "구체 행동"], ch9["tasks"], [1, 1.6, 3.4])
    b.h2("3. 한계 및 유의사항")
    for l in ch9["limitations"]:
        b.bullet(l)
    if ch9.get("evidence_grade"):
        b.h2("4. 분석영역별 근거 등급")
        b.table(["분석영역", "근거등급", "현재 근거", "추가 검증 방법"], ch9["evidence_grade"], [1, 0.8, 2.2, 2.5])
    b.note("공백기술 및 IP 전략과 관련한 실제 출원 여부 판단은 반드시 전문 변리사의 선행기술조사(FTO) 및 정밀 검토를 거쳐야 한다.")

    return b.save_to_buffer()
